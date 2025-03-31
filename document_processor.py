from langchain_community.document_loaders import PyPDFLoader, UnstructuredFileLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
import logging

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, chunk_size=1000, chunk_overlap=100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

    def load_documents(self, directory_path):
        documents = []
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)
            if filename.endswith('.pdf'):
                loader = PyPDFLoader(file_path)
                docs = loader.load()
            else:
                loader = UnstructuredFileLoader(file_path)
                docs = loader.load()
            
            for doc in docs:
                doc.metadata['source'] = filename
                doc.metadata['file_path'] = file_path
                doc.metadata['original_filename'] = filename  # Add original filename
            
            documents.extend(docs)
        return documents

    def chunk_documents(self, documents):
        chunks = self.text_splitter.split_documents(documents)
        for i, chunk in enumerate(chunks):
            chunk.metadata['chunk_id'] = i
            
            # Ensure original filename is preserved in chunks
            if 'original_filename' not in chunk.metadata and 'source' in chunk.metadata:
                chunk.metadata['original_filename'] = chunk.metadata['source']
                
        return chunks
        
    def process_file(self, file_path, original_filename=None):
        try:
            # Initialize documents variable
            documents = []
            
            if file_path.endswith('.pdf'):
                loader = PyPDFLoader(file_path)
                documents = loader.load()
            elif file_path.endswith('.docx'):
                try:
                    from docx import Document
                    doc = Document(file_path)
                    # Extract text from all paragraphs
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    # Create a document with the extracted text
                    from langchain_core.documents import Document as LangchainDocument
                    documents = [LangchainDocument(page_content=text)]
                except ImportError:
                    raise ImportError(
                        "To process DOCX files, please install the required dependencies:\n"
                        "pip install python-docx"
                    )
            else:
                loader = UnstructuredFileLoader(file_path)
                documents = loader.load()
                
            filename = os.path.basename(file_path)
            
            # Always use the original filename if provided, otherwise use the current filename
            display_filename = original_filename if original_filename else filename
            
            # Log the raw document content
            logger.info(f"Processing file: {display_filename}")
            logger.info(f"File path: {file_path}")
            
            # Use the provided original filename if available
            for doc in documents:
                # Set metadata consistently
                doc.metadata['source'] = display_filename  # Use display filename for source
                doc.metadata['file_path'] = file_path
                doc.metadata['original_filename'] = display_filename  # Store original filename
                doc.metadata['temp_filename'] = filename  # Store temp filename for internal use
                
                # Log document details
                logger.info(f"Document metadata: {doc.metadata}")
                logger.info(f"Document length: {len(doc.page_content)} characters")
                logger.info(f"Document preview: {doc.page_content[:500]}...")
            
            chunks = self.chunk_documents(documents)
            
            # Log the chunked content
            logger.info(f"Generated {len(chunks)} chunks for {display_filename}")
            for i, chunk in enumerate(chunks):
                # Ensure metadata is preserved in chunks
                chunk.metadata['source'] = display_filename
                chunk.metadata['original_filename'] = display_filename
                chunk.metadata['chunk_id'] = i
                
                logger.info(f"Chunk {i} metadata: {chunk.metadata}")
                logger.info(f"Chunk {i} length: {len(chunk.page_content)} characters")
                logger.info(f"Chunk {i} preview: {chunk.page_content[:500]}...")
            
            return chunks
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            raise